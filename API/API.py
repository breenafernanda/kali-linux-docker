import os, re, time, requests, docx, cv2, pytesseract, utm
import numpy as np
# import pyautogui as pg
from flask import Flask, request, jsonify
from docx.shared import Inches
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from PIL import Image, ImageTk,ImageEnhance
from datetime import datetime
import pandas as pd
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By

app = Flask(__name__)


@app.route('/memorial_descritivo', methods=['POST'])
async def demorial_descritivo():

    receive_data = request.json
   
    """_________________________________________________________________________________"""

    def salvar_memorial_descritivo(cliente):
        # imprimir('\n')
        
        try:
            document.save(f'Arquivos\Memorial Descritivo_{cliente}.docx')
            print(f'Memorial Descritivo_{cliente}.docx')

        except: 
            # imprimir(f'Erro ao salvar, fechar arquivo e tentar novamente')
            try:
                i = 1
                aberto = True
                while (aberto==True):
                    try:
                        document.save(f'Arquivos\Memorial Descritivo_{cliente}({i}).docx')
                        print(f'Memorial Descritivo_{cliente}({i}).docx')
                        aberto = '✘'
        #                 i=0
                    except:
                        i = i + 1
            except Exception as e: print(e)
    
    def salvar_formulario_microGD(cliente):
        # imprimir('\n')
        # Salvar o arquivo modificado
        try:
            workbook.save(f'Arquivos\Formulario_MicroGD_{cliente}.xlsx')
            print(f'Formulario_MicroGD_{cliente}.xlsx')

        except: 
            # imprimir(f'Erro ao salvar, fechar arquivo e tentar novamente')
            try:
                i = 1
                aberto = True
                while (aberto==True):
                    try:
                        workbook.save(f'Arquivos\Formulario_MicroGD_{cliente}({i}).xlsx')
                        print(f'Formulario_MicroGD_{cliente}({i}).xlsx')
                        aberto = '✘'
        #                 i=0
                    except:
                        i = i + 1
            except Exception as e: print(e)

    def adicionar_dados_cadastro(dado):
        # obter célula de Nome
        if dado == 'cliente':
            table = document.tables[0] 
            cell = table.cell(0,0) # índices começando em 0

        if dado == 'data':
            table = document.tables[1] # assume que a tabela desejada é a primeira na ordem de aparição
            cell = table.cell(0,0)

        if dado == 'endereço':
            table = document.tables[2] # assume que a tabela desejada é a primeira na ordem de aparição
            cell = table.cell(0,0)

        if dado == 'endereço_print':
            table = document.tables[2]
            cell = table.cell(0,1)

        if dado == 'descricao_geral_II':
            table = document.tables[3] # assume que a tabela desejada é a primeira na ordem de aparição
            cell = table.cell(0,0)

        if dado == 'descri_geraçao_I':
            table = document.tables[4]
            cell = table.cell(0,0)

        if dado == 'titulo_3_1':
            table = document.tables[5]
            cell = table.cell(0,0)
        
        if dado == 'caracteristicas_inversor':
            table = document.tables[5]
            cell = table.cell(0,1)
        
        if dado == 'titulo_3_2':
            table = document.tables[6]
            cell = table.cell(0,0)
        
        if dado == 'caracteristicas_modulo':
            table = document.tables[6]
            cell = table.cell(0,1)
        
        if dado == 'previsao':
            table = document.tables[7]
            cell = table.cell(0,0)
        return table, cell
    
    def inserir_dados(dado, texto):
        # inserir nome 
        try:
            table, cell = adicionar_dados_cadastro(dado)
            if texto == './screenshot.png':
                # inserir imagem na celula
                cell_paragraph = cell.paragraphs[0]
                run = cell_paragraph.add_run()
                run.add_picture(texto, width=Inches(8.0), height=Inches(4.0))
            else:    
                cell.text = texto

                # Modificar estilo da célula
                run = cell.paragraphs[0].runs[0]
                run.font.name = 'Arial'  # Define a fonte como Arial

                if dado == 'cliente':
                    run.font.size = Pt(26)  # Ajuste o tamanho da fonte para 26 pontos
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Ajuste o alinhamento conforme necessário
                elif dado == 'titulo_3_1' or dado == 'titulo_3_2':
                    run.font.size = Pt(12)  # Ajuste o tamanho da fonte para 26 pontos
                    run.font.bold = True
                else:
                    run.font.size = Pt(12)  # Ajuste o tamanho da fonte para 26 pontos

                # run.font.bold = True  # Define o texto em negrito
        except: 
            print(f'ERRO ao inserir {dado} - > {texto}')

    def obter_dia_hoje():
        # Obter a data atual
        data_atual = datetime.now()

        # Mapear os nomes dos meses
        meses = [
            "janeiro", "fevereiro", "março", "abril", "maio", "junho",
            "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"
        ]

        # Formatar a data no estilo desejado
        data_formatada = "{} de {} de {}".format(
            data_atual.day, meses[data_atual.month - 1], data_atual.year
        )

        return data_formatada

    def obter_dados_inversores(modelo):
            caminho_do_arquivo = './config/inversores.xlsx'

            # Ler arquivo Excel
            dados_excel = pd.read_excel(caminho_do_arquivo)

            # Iterar sobre as linhas do DataFrame
            for indice, linha in dados_excel.iterrows():
                # Converter a linha em uma lista
                valores_da_linha = linha.tolist()
                if valores_da_linha[1] == modelo:
                    # Faça algo com a lista de valores, por exemplo, imprimir
                    print(f'Valores da Linha {indice}: {valores_da_linha}')
                    return valores_da_linha
    
    def obter_dados_modulos(modelo):
            caminho_do_arquivo = './config/modulos.xlsx'

            # Ler arquivo Excel
            dados_excel = pd.read_excel(caminho_do_arquivo)

            # Iterar sobre as linhas do DataFrame
            for indice, linha in dados_excel.iterrows():
                # Converter a linha em uma lista
                valores_da_linha = linha.tolist()
                if valores_da_linha[1] == modelo:
                    # Faça algo com a lista de valores, por exemplo, imprimir
                    print(f'Valores da Linha {indice}: {valores_da_linha}')
                    return valores_da_linha

    def extrair_print_google_maps(url):
        def converter_coord_em_utm(coordenadas):
            # Extrai os valores de graus, minutos e segundos da string de coordenadas
            match = re.match(r'(\d+)°(\d+)\'([\d.]+)"([NS]) (\d+)°(\d+)\'([\d.]+)"([WE])', coordenadas)
            
            if not match:
                raise ValueError("Formato de coordenadas inválido")

            # Extrai os valores da correspondência
            lat_graus, lat_min, lat_sec, lat_dir, lon_graus, lon_min, lon_sec, lon_dir = match.groups()

            # Converte os valores para números decimais
            lat_decimal = float(lat_graus) + float(lat_min)/60 + float(lat_sec)/3600
            lon_decimal = float(lon_graus) + float(lon_min)/60 + float(lon_sec)/3600

            # Determina o sinal da latitude e longitude
            if lat_dir.upper() == 'S':
                lat_decimal = -lat_decimal
            if lon_dir.upper() == 'W':
                lon_decimal = -lon_decimal

            # Converte para UTM
            utm_coords = utm.from_latlon(lat_decimal, lon_decimal)

            # Extrai os valores UTM
            abcissa, ordenada, fuso, hemisferio = utm_coords

            return int(fuso), int(ordenada), int(abcissa), hemisferio

        try:
            # Define o diretório do perfil do Instagram
            profile = 'API'
            dir_path = os.getcwd()
            profile = os.path.join(dir_path, profile)
            
            options = webdriver.ChromeOptions()
            
            options.add_argument(r"user-data-dir={}".format(profile))
            options.add_argument('--disable-gpu')
            options.add_argument('--no-sandbox')   
            options.add_argument('--headless')
            options.use_chromium = True
            options.add_experimental_option('excludeSwitches', ['enable-logging'])
            options.add_experimental_option('prefs', {'download.default_directory': os.path.join(os.getcwd(), 'geracao')})
            options.add_argument("--window-size=1920x1080")
            try:
                driver_manager = ChromeDriverManager()
                driver = webdriver.Chrome(executable_path=driver_manager.install(), options=options)
            except Exception as e: 
                print('........',e)
                driver_manager = ChromeDriverManager().install()
                driver = webdriver.Chrome(options=options)

            print(f'\n💻 \x1b[32m Navegador Chrome -({profile}) iniciado! \x1b[0m✅\n')
            driver.maximize_window()
            
            # acessar url
            driver.get(url)
            # Executar script JavaScript para ajustar o zoom
            script = f"document.body.style.zoom='140%'"
            driver.execute_script(script)
            # Salvar a captura de tela
            driver.save_screenshot('./screenshot.png')
            print(f"Captura de tela salva em: ./screenshot.png")



            # Extrair dados de latitude e longitude
            coordenadas = driver.find_element(By.CSS_SELECTOR, '#QA0Szd > div > div > div.w6VYqd > div:nth-child(2) > div > div.e07Vkf.kA9KIf > div > div > div.TIHn2 > div > div.lMbq3e > div:nth-child(1) > h1').text
            print("Coordenadas obtidas:", coordenadas)
            fuso, ordenada, abcissa, hemisferio = converter_coord_em_utm(coordenadas)

            # Imprime as coordenadas UTM separadamente
            print("Fuso:", fuso)
            print("Ordenada:", ordenada)
            print("Abcissa:", abcissa)
            print("Hemisfério:", hemisferio)
            driver.close()
            driver.quit()
            return fuso, ordenada, abcissa
        except Exception as erro: print(f'🚨 : {erro}')
    

    print(f'\x1b[34m************************************************\nDados recebidos pela API:\n{receive_data}\n************************************************\n\n\x1b[0m')


    response = None
    # EXECUTAR PROCEDIMENTOS MEMORIAL DESCRITIVO
        # __________________________________________________________
    # MEMORIAL DESCRITIVO 
    # __________________________________________________________

    print(f'\x1b[36mEXECUTANDO API PROJETOS -  Memorial Descritivo\n\x1b[0m')
    try:
        print(receive_data)
        try:
            numero_proposta = receive_data['numero_proposta']
            cliente = receive_data['cliente']
            previsao_mensal = receive_data['previsao_mensal']
            google_maps = receive_data['google_maps']
            fuso, ordenada, abcissa = extrair_print_google_maps(google_maps)
            response = '200 - Dados recebidos com sucesso!' ## setpoint
        except Exception as e: print(f'Erro ao extrair os dados da API: {e}')
        
        ## manipular arquivo  Memorial Descritivo
        try:
            ## abrir documento Memorial descritivo.docx
            document = docx.Document(r'Memorial descritivo.docx')

            # Inserir Nome da usina na primeira página
            inserir_dados('cliente', cliente)

            # Inserir data de hoje
            data_hoje = obter_dia_hoje()
            inserir_dados('data', data_hoje)

            # 2. DESCRIÇÃO GERAL DO CONSUMIDOR
            ## endereço
            end_rua = receive_data['end_rua']
            end_numero = receive_data['end_numero']
            end_bairro = receive_data['end_bairro']
            end_CEP = receive_data['end_CEP']
            end_cidade = receive_data['end_cidade']
            end_UF = receive_data['end_UF']
            print(end_rua)
            endereco_formatado = f'{end_rua} nº {end_numero} - {end_bairro} - CEP: {end_CEP}   {end_cidade}-{end_UF}'

            inserir_dados('endereço', endereco_formatado)
            print(f'\x1b[32mENDEREÇO: {endereco_formatado}\x1b[0m')


            # INSERIR PRINT DO GOOGLE MAPS
            inserir_dados('endereço_print', './screenshot.png')


            tipo_conexao = receive_data['tipo_conexão']
            disjuntor_entrada = receive_data['disjuntor_amperagem']
            tipo_conexao = receive_data['tipo_conexão']
            id_cliente = receive_data['id_cliente']
            codigo_instalação = receive_data['codigo_instalação']
            descricao_geral_II = (
                    f'  - {tipo_conexao};\n'
                    f'  - {disjuntor_entrada};\n'
                    f'  - Cliente: {id_cliente};\n'
                    f'  - Instalação: {codigo_instalação};'
            )

            inserir_dados('descricao_geral_II', descricao_geral_II)


            # 3. DESCRIÇÃO GERAL DA GERAÇÃO DISTRIBUÍDA
            qtd_inversor = receive_data['qtde_inversor']
            potencia_inversor = receive_data['potencia_inversor']
            modelo_inversor = receive_data['id_inversor']

            qtd_modulos = receive_data['qtde_modulo']
            potencia_modulos = receive_data['potencia_modulo']
            modelo_modulos = receive_data['id_modulo']

            if qtd_inversor == 1:
                texto_descricao_geral_da_geracao_distribuida = (
                            f'Gerador Fotovoltaico Grid-Tie composto de {qtd_inversor} inversor fotovoltaico Solar {potencia_inversor} W On-Grid de modelo {modelo_inversor}, e {qtd_modulos} painéis de {potencia_modulos} W fotovoltaicos de modelo {modelo_modulos}.'
                            
                                                            )
            else:
                texto_descricao_geral_da_geracao_distribuida = (
                            f'Gerador Fotovoltaico Grid-Tie composto de {qtd_inversor} inversores fotovoltaico Solar {potencia_inversor} W On-Grid de modelo {modelo_inversor}, e {qtd_modulos} painéis de {potencia_modulos} W fotovoltaicos de modelo {modelo_modulos}.'
                            
                                                            )  
            inserir_dados('descri_geraçao_I',texto_descricao_geral_da_geracao_distribuida)
            print(f'\x1b[33mDESCRIÇÃO GERAL DA GERAÇÃO DISTRIBUÍDA:\n {texto_descricao_geral_da_geracao_distribuida}\x1b[0m')

            # DESCRIÇÃO DO INVERSOR
            titulo_3_1 = f'3.1   Inversor Solar {potencia_inversor}W {modelo_inversor};\n'
            
            valores = obter_dados_inversores(modelo_inversor)
            if valores[2] != '***':
                print(f'INMETRO : {valores[2]}')
                titulo_3_1 += f'        INMETRO: {valores[2]};\n'
            inserir_dados('titulo_3_1', titulo_3_1)
            marca_inversor = valores[0]

            caracteristicas_inversor = (
                f'Características da Entrada:\n'
                f'•	Tensão de entrada máxima: {valores[3]} V \n'
                f'•	Faixa de tensão de operação do MPPT: {valores[4]}\n'
                f'•	Tensão de partida: {valores[5]} V \n'
                f'•	Corrente de entrada máxima por MPPT: {valores[6]} A \n'
                f'•	Corrente de curto-circuito máxima: {valores[7]} A \n'
                f'•	Número de MPPTs: {valores[8]}\n'
                f'•	Número máximo de entradas por MPPT: {valores[9]} \n\n'
                f'Características da Saída:\n'
                f'•	Conexão à rede {valores[10]} \n'
                f'•	Potência nominal de saída: {valores[11]} W\n'
                f'•	Potência aparente máxima: {valores[12]} VA\n'
                f'•	Tensão de saída nominal: {valores[13]} V \n'
                f'•	Frequência de rede CA nominal: {valores[14]} \n'
                f'•	Corrente de saída máxima: {valores[15]} A \n'
                f'•	Fator de potência ajustável: {valores[16]} \n'
                f'•	Distorção harmônica total máxima:  ≤ {valores[17]} %\n\n'
                f'Dados Gerais:\n'
                f'•	Faixa de temperatura de operação: {valores[18]} \n'
                f'•	Umidade relativa de operação: {valores[19]} \n'
                f'•	Altitude de operação: {valores[20]} \n'
                f'•	Resfriamento {valores[21]} \n'
                f'•	Display {valores[22]} \n'
                f'•	Comunicação {valores[23]} \n'
                f'•	Peso: (incluindo suporte de montagem) {valores[24]} kg \n'
                f'•	Dimensão: (incluindo suporte de montagem) {valores[25]} mm \n'
                f'•	Grau de proteção: {valores[26]}\n'
            )

            inserir_dados('caracteristicas_inversor', caracteristicas_inversor)
            
            # DESCRIÇÃO DOS MÓDULOS
            titulo_3_2 = f'3.2  Painel Solar {modelo_modulos}'

            inserir_dados('titulo_3_2', titulo_3_2)

            valores = obter_dados_modulos(modelo_modulos)

            print(valores)
            marca_modulo = valores[0]

            caracteristicas_modulo = (
                f'•	Potência máxima (Pmax): {valores[2]}Wp\n'
                f'•	Tensão em circuito aberto (Voc): {valores[3]} V\n'
                f'•	Tensão de Pico (Vmpp): {valores[4]} V\n'
                f'•	Corrente de curto-circuito (Isc): {valores[5]} A\n'
                f'•	Corrente de Pico (Impp): {valores[6]} A\n'
                f'•	Tipo de célula: Silício {valores[7]}\n'
                f'•	Dimensões painel:{valores[8]} (mm)\n'
            )

            inserir_dados('caracteristicas_modulo',caracteristicas_modulo)


            # 4 . PREVISÃO DA PRODUÇÃO DE ENERGIA
            texto_previsao = (
                f'A Estimativa de geração mensal é de {previsao_mensal} kWh. O Sistema funciona aproximadamente 12 horas diárias, das 06h às 18h.'
            )
            inserir_dados('previsao', texto_previsao)


            #salvar arquivo
            salvar_memorial_descritivo(cliente)
            response = '200 - Arquivo editado com sucesso!'
        except Exception as e: print(f'Erro ao salvar o arquivo: {e}')
    
    except Exception as e: print(e)


    # __________________________________________________________
    # FORMULÁRIO MICROGD
    # __________________________________________________________

    def obter_cell_formulario(dado):
        if dado == 'num_cliente':
            table = document.tables[1] 
            cell = table.cell(0,1) # índices começando em 0
       
        if dado == 'id_instalação':
            table = document.tables[1] 
            cell = table.cell(0,3) # índices começando em 0
        
        if dado == 'titular_unidade':
            table = document.tables[2] 
            cell = table.cell(0,1) # índices começando em 0
        
        if dado == 'grupo':
            table = document.tables[3] 
            cell = table.cell(0,1) # índices começando em 0
           
        if dado == 'classe':
            table = document.tables[3] 
            cell = table.cell(0,3) # índices começando em 0
        
        if dado == 'documento':
            table = document.tables[3] 
            cell = table.cell(0,5) # índices começando em 0
           
        if dado == 'logradouro':
            table = document.tables[4] 
            cell = table.cell(0,1) # índices começando em 0
           
        if dado == 'end_numero':
            table = document.tables[4] 
            cell = table.cell(0,3) # índices começando em 0
        
        if dado == 'end_complemento':
            table = document.tables[4] 
            cell = table.cell(0,5) # índices começando em 0
        
        if dado == 'bairro':
            table = document.tables[5] 
            cell = table.cell(0,1) # índices começando em 0
           
        if dado == 'cidade':
            table = document.tables[5] 
            cell = table.cell(0,3) # índices começando em 0
           
        if dado == 'estado':
            table = document.tables[5] 
            cell = table.cell(0,5) # índices começando em 0

        if dado == 'CEP':
            table = document.tables[5] 
            cell = table.cell(0,7) # índices começando em 0
        
        if dado == 'telefone':
            table = document.tables[6] 
            cell = table.cell(0,1) # índices começando em 0
           
        if dado == 'celular':
            table = document.tables[6] 
            cell = table.cell(0,3) # índices começando em 0

        if dado == 'email':
            table = document.tables[6] 
            cell = table.cell(0,5) # índices começando em 0
        
        """ 2 %%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%"""
        # if dado == 'fuso':
        #     table = document.tables[8]
        #     cell = table.cell(0,1)

        # if dado == 'abscissa':
        #     table = document.tables[8]
        #     cell = table.cell(0,3)

        # if dado == 'ordenada':
        #     table = document.tables[8]
        #     cell = table.cell(0,5)

        # if dado == 'consumo_proprio':
        #     table = document.tables[9]
        #     cell = table.cell(0,1)

        # if dado == 'potencia_motor_gerador':
        #     table = document.tables[10]
        #     cell = table.cell(0,1)
        
        # if dado == 'tipo_de_subestação':
        #     table = document.tables[11]
        #     cell = table.cell(0,2)
        
        # if dado == 'transformador_particular':
        #     table = document.tables[11]
        #     cell = table.cell(1,2)
        
        # if dado == 'tipo_de_solicitação':
        #     table = document.tables[12]
        #     cell = table.cell(0,1)
        
        # if dado == 'tipo_de_edificação':
        #     table = document.tables[13]
        #     cell = table.cell(0,1)

        # if dado == 'disjuntor_tipo':
        #     table = document.tables[14]
        #     cell = table.cell(0,2)
        
        # if dado == 'disjuntor_amperagem':
        #     table = document.tables[14]
        #     cell = table.cell(0,3)

        # if dado == 'tensao_atendimento':
        #     table = document.tables[15]
        #     cell = table.cell(0,1)
        
        # if dado == 'tipo_de_ramal':
        #     table = document.tables[15]
        #     cell = table.cell(0,3)
        
        # if dado == 'mudança_padrão_entrada':
        #     table = document.tables[16]
        #     cell = table.cell(0,1)

        """ 4 %%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%"""
        
        # if dado == 'tipo_fonte_primaria':
        #     table = document.tables[19]
        #     cell = table.cell(0,1)
        
        # if dado == 'potencia_ativa_total':
        #     table = document.tables[19]
        #     cell = table.cell(0,3)
        
        # if dado == 'tipo_de_geração':
        #     table = document.tables[20]
        #     cell = table.cell(0,1)
        
        # if dado == 'modalidade_compensação':
        #     table = document.tables[21]
        #     cell = table.cell(0,1)
        
        # if dado == 'qtde_compensação':
        #     table = document.tables[21]
        #     cell = table.cell(0,3)
        
        # if dado == 'potencia_total_modulos':
        #     table = document.tables[23]
        #     cell = table.cell(0,1)
        
        # if dado == 'potencia_total_inversores':
        #     table = document.tables[23]
        #     cell = table.cell(1,1)
        
        # if dado == 'area_arranjos':
        #     table = document.tables[23]
        #     cell = table.cell(2,1)

        # if dado == 'qtdade_modulos':
        #     table = document.tables[23]
        #     cell = table.cell(3,1)
        
        # if dado == 'modelo_modulo':
        #     table = document.tables[23]
        #     cell = table.cell(4,1)
        
        # if dado == 'marca_modulo':
        #     table = document.tables[23]
        #     cell = table.cell(5,1)
            
        # if dado == 'qtde_inversor':
        #     table = document.tables[23]
        #     cell = table.cell(6,1)
        
        # if dado == 'modelo_inversor':
        #     table = document.tables[23]
        #     cell = table.cell(7,1)
        
        # if dado == 'marca_inversor':
        #     table = document.tables[23]
        #     cell = table.cell(8,1)
        
        # if dado == 'procurador_legal':
        #     table = document.tables[33]
        #     cell = table.cell(0,1)
        
        # if dado == 'end_procurador':
        #     table = document.tables[34]
        #     cell = table.cell(0,1)
        
        # if dado == 'telefone_procurador':
        #     table = document.tables[35]
        #     cell = table.cell(0,1)
        
        # if dado == 'celular_procurador':
        #     table = document.tables[35]
        #     cell = table.cell(0,3)

        # if dado == 'email_procurador':
        #     table = document.tables[35]
        #     cell = table.cell(0,5)

        # if dado == 'local_data':
        #     table = document.tables[37]
        #     cell = table.cell(2,0)

        return table, cell
    

    def editar_formulario_word(dado, texto):
            table, cell = obter_cell_formulario(dado)
            cell.text = texto
            # Modificar estilo da célula
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Calibri'  # Define a fonte como Arial
            run.font.size = Pt(11)  # Ajuste o tamanho da fonte
            # run.font.bold = True # NEGRITO
            run.font.color.rgb = RGBColor(11, 156, 245)  # Define a cor como azul claro (RGB)
    def editar_formulario(sheet, celula, texto):
        # Selecionar a opção desejada no campo 'I39'
        sheet[celula].value = texto
    ## manipular arquivo

    # 1 - IDENTIFICAÇÃO DA UNIDADE CONSUMIDORA - UC
    try:
        import openpyxl
        from openpyxl import load_workbook
        from openpyxl.worksheet.datavalidation import DataValidation

        # Caminho para o arquivo Excel
        arquivo_excel = 'Formulario-MicroGD.xlsx'

        # Carregar o arquivo Excel
        workbook = load_workbook(arquivo_excel)

        # Nome da planilha
        nome_planilha = 'Formulário'  # Certifique-se de ajustar para o nome real da planilha

        # Selecionar a planilha desejada
        sheet = workbook[nome_planilha]

        sheet['I12'].value = id_cliente # Numero do Cliente
        sheet['AG12'].value = codigo_instalação # Número da Instalação
        sheet['N14'].value = cliente # Titular da Unidade Consumidora
        sheet['E16'].value =  receive_data['grupo_conexão'] # Grupo
        sheet['K16'].value =  receive_data['classe_conexão'] # Classe
        sheet['AC16'].value =  receive_data['cliente_CPF_CNPJ'] # CPF/CNPJ 
        sheet['G18'].value = receive_data['end_rua'] # Logradouro
        sheet['AI18'].value = receive_data['end_numero'] # Número
        sheet['AR18'].value = receive_data['end_complemento'] # Complemento
        sheet['E20'].value = receive_data['end_bairro'] # Bairro
        sheet['T20'].value = receive_data['end_cidade'] # Município
        sheet['AN20'].value = receive_data['end_UF'] # Estado
        sheet['AS20'].value = receive_data['end_CEP'] # CEP
        sheet['F22'].value = receive_data['cliente_telefone'] # Telefone
        sheet['O22'].value = receive_data['cliente_celular'] # Celular
        sheet['Y22'].value =  receive_data['cliente_email'] # E-mail
        # sheet[''].value =  # 
    except Exception as e: print(e)

    # # 2 - DADOS DA UNIDADE CONSUMIDORA %%% REFORMULAR PARA RECEBER EM DATA ESSAS INFO
    try:
        # LOCALIZAÇÃO: FUSO, ABSCISSA E ORDENADA
        sheet['V27'].value = fuso
        sheet['AC27'].value = abcissa
        sheet['AL27'].value = ordenada
        sheet['I39'].value = receive_data['tipo_solicitação']
        sheet['I41'].value = receive_data['tipo_edificação']

        sheet['AB45'].value = receive_data['disjuntor_tipo']
        sheet['AH45'].value = receive_data['disjuntor_amperagem']
        sheet['L52'].value = receive_data['tensao_atendimento']
        sheet['X52'].value = receive_data['ramal_instalação']
        sheet['R54'].value = receive_data['mudança_padrão_entrada']
  
    except Exception as e: print(f'\x1b[36m>>> Erro em # 2 - Dados da Unidade Consumidora : \n {e}\x1b[0m')
    
    # # 4 - DADOS DA GERAÇÃO
    try:
        sheet['L84'].value = receive_data['modalidade_compensação']
        sheet['AS84'].value = receive_data['qtd_compensação']

        potencia_modulo = int(receive_data['potencia_modulo'])
        qtdade_modulos = int(receive_data['qtde_modulo'])
        potencia_total_modulos = (potencia_modulo * qtdade_modulos)/1000
        sheet['P92'].value = potencia_total_modulos
        potencia_inversor = int(receive_data['potencia_inversor'])
        qtdade_inversor = int(receive_data['qtde_inversor'])
        potencia_total_inversores = (potencia_inversor * qtdade_inversor)/1000
        sheet['P94'].value = potencia_total_inversores
        
        sheet['P96'].value = receive_data['area_arranjos']
        sheet['P98'].value = receive_data['qtde_modulo']
        sheet['P100'].value = receive_data['id_modulo']
        sheet['P102'].value = marca_modulo
        sheet['P104'].value = receive_data['qtde_inversor']
        sheet['P106'].value = receive_data['id_inversor']
        sheet['P108'].value = marca_inversor

    
        # editar_formulario('area_arranjos', receive_data['area_arranjos'])
        # editar_formulario('qtdade_modulos', receive_data['qtde_modulo'])
        # editar_formulario('modelo_modulo',receive_data['id_modulo'])
        # editar_formulario('marca_modulo', marca_modulo)
        # editar_formulario('qtde_inversor', receive_data['qtde_inversor'])
        # editar_formulario('modelo_inversor', receive_data['id_inversor'])
        # editar_formulario('marca_inversor', marca_inversor)

    except Exception as e: print(f'\x1b[36m>>> Erro em # 4 - Dados da Geração: \n {e}\x1b[0m')

    # # 9 - SOLICITANTE (DADOS ENGENHEIRO RESPONSÁVEL)
    try:

        dados_livimar = {
            'procurador_legal':'LIVIMAR PINHEIRO DE OLIVEIRA JUNIOR',
            'endereço': endereco_formatado,     # confirmar se serão dados do Livimar ou do cliente
            'telefone':receive_data['cliente_telefone'], 
            'celular':receive_data['cliente_celular'],
            'email': receive_data['cliente_email']
        }

        # inserir dados do Livimar (Engenheiro)
        sheet['O196'].value = dados_livimar['endereço']
        sheet['F200'].value = dados_livimar['telefone']
        sheet['O200'].value = dados_livimar['celular']
        sheet['Y200'].value = dados_livimar['email']
        sheet['C210'].value =  f'UBERABA-MG, {data_hoje}'
        # editar_formulario('procurador_legal', dados_livimar['procurador_legal'])
        # editar_formulario('end_procurador', dados_livimar['endereço'])
        # editar_formulario('telefone_procurador', dados_livimar['telefone']) 
        # editar_formulario('celular_procurador', dados_livimar['celular'])
        # editar_formulario('email_procurador', dados_livimar['email'])
        # editar_formulario('local_data', f'Uberaba-MG, {data_hoje}')

    except Exception as e: print(f'\x1b[36m>>> Erro em # 9 - Solicitante: \n {e}\x1b[0m') 


    # salvar formulário
    try:
        salvar_formulario_microGD(cliente)
    except Exception as e: print(f'\x1b[36m>>> Erro em # SALVAR ARQUIVO DOCX: \n {e}\x1b[0m')


    # converter docx para pdf

    return jsonify(response)



if __name__ == '__main__':

    app.run(port=5000, host='0.0.0.0', debug=True)

